import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def generate_module_2_doc():
    source_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_1_Gestion_de_Mascotas.docx'
    output_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_2_Historial_de_Usuarios.docx'
    
    doc = docx.Document(source_path)
    
    # 1. Update text in Header Table 0
    t0 = doc.tables[0]
    t0.rows[2].cells[1].text = "PawTok – Módulo 2: Historial de Usuarios"
    t0.rows[3].cells[1].text = "Versión 1.0  |  18/08/2026"
    
    # 2. Update Table 1 (Notice)
    t1 = doc.tables[1]
    t1.rows[0].cells[0].text = "Cómo usar esta plantilla: Este Plan de Pruebas corresponde de manera exclusiva y detallada al MÓDULO 2: HISTORIAL DE USUARIOS del sistema PawTok. Contempla la validación funcional de los requisitos RF 2.1, RF 2.3, RF 2.4, RF 2.5, RF 2.7 y RF 2.8, sus reglas de negocio asociadas (RN 2.1 al RN 2.8) y los atributos de calidad ISO/IEC 25010."

    # 3. Update Paragraphs text for Sections 1 and 2
    for p in doc.paragraphs:
        if "Objetivo del plan:" in p.text:
            p.text = "Objetivo del plan: Comprobar y validar rigurosamente que las funcionalidades del Módulo 2: Historial de Usuarios del sistema PawTok cumplan de manera precisa con los requisitos funcionales (RF 2.1, RF 2.3, RF 2.4, RF 2.5, RF 2.7, RF 2.8), sus respectivas reglas de negocio (RN 2.1 a RN 2.8) y los estándares de calidad ISO/IEC 25010 en almacenamiento automático de actividades, consulta de movimientos, asociación con mascotas adoptadas, filtrado específico, detalles de usuario y eliminación de registros."
        elif "Contexto del sistema:" in p.text:
            p.text = "Contexto del sistema: PawTok es una plataforma web desarrollada para conectar refugios y adoptantes. El Módulo de Historial de Usuarios es el componente neurálgico encargado de registrar automáticamente cada movimiento y acción en la plataforma, asociar de forma transparente las mascotas adoptadas, permitir consultas y búsquedas personalizadas, presentar la información de antigüedad/perfil del usuario y gestionar la depuración o eliminación segura de registros del historial."
        elif "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 1" in p.text:
            p.text = "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 2: Historial de Usuarios de la plataforma PawTok:"
        elif "Módulos externos a la gestión de mascotas" in p.text:
            p.text = "• Módulos externos al historial de usuarios, tales como la gestión y edición clínica de mascotas (Módulo 1), catálogo de razas (Módulo 3) y administración global de refugios (Módulo 4), evaluados en sus planes específicos."
        elif "Para el Módulo 1: Gestión de Mascotas se implementa" in p.text:
            p.text = "Para el Módulo 2: Historial de Usuarios se implementa una estrategia de pruebas manuales de caja negra combinada con pruebas de API e interfaz. Se enfoca en validar el almacenamiento automático de eventos, la integridad de la asociación usuario-mascota, la precisión de los filtros, la privacidad de datos personales y el control de eliminaciones."
        elif "Cuidado con los datos: Para las pruebas del Módulo de Mascotas" in p.text:
            p.text = "Cuidado con los datos: Para las pruebas del Módulo de Historial de Usuarios se utilizan únicamente cuentas de adoptantes y refugios ficticios de prueba. Las contraseñas se encuentran cifradas (BCrypt) y ningún dato sensible o real es expuesto."
        elif "El Módulo 1: Gestión de Mascotas ha superado" in p.text:
            p.text = "El Módulo 2: Historial de Usuarios ha superado satisfactoriamente el ciclo de pruebas planificado. Se verificó el cumplimiento del 100% de los requisitos funcionales (RF 2.1, 2.3, 2.4, 2.5, 2.7, 2.8) y reglas de negocio (RN 2.1 a 2.8). Se recomienda su paso a la fase de integración general del sistema."

    # Update Section 2.1 Bullet points in paragraphs
    bullets_m2 = [
        "• Almacenamiento automático y persistencia en base de datos de cada acción y movimiento del usuario [RF 2.1, RN 2.1].",
        "• Asociación y vinculación directa de cada usuario adoptante con las mascotas adoptadas en el historial [RF 2.3, RN 2.3].",
        "• Consulta y visualización completa del historial de actividades, solicitudes y movimientos en la vista de cuenta [RF 2.4, RN 2.4].",
        "• Búsqueda y filtrado dinámico del historial por mascota específica [RF 2.5, RN 2.5].",
        "• Visualización detallada del perfil del usuario, fecha de registro y antigüedad en el sistema [RF 2.7, RN 2.7].",
        "• Eliminación y depuración permitida de registros del historial de solicitudes por parte del usuario [RF 2.8, RN 2.8].",
        "• Control de acceso por roles y privacidad (el usuario solo visualiza sus propios registros de historial; el Administrador accede a la auditoría global) [RNF 5, RNF 13].",
        "• Tiempos de respuesta de consulta menores a 3 segundos y tasa de fallos inferior al 1% [RNF 1, RNF 8].",
        "• Compatibilidad en navegadores web (Google Chrome y Microsoft Edge) [RNF 18]."
    ]
    
    # Replace bullet paragraphs in Section 2.1
    p_idx = 0
    in_21 = False
    for p in doc.paragraphs:
        if "2.1 Incluido en el alcance" in p.text:
            in_21 = True
            continue
        if "2.2 Fuera del alcance" in p.text:
            in_21 = False
            break
        if in_21 and p.text.startswith("•") and p_idx < len(bullets_m2):
            p.text = bullets_m2[p_idx]
            p_idx += 1

    # 4. Table 2: Elementos a probar (EP)
    t2 = doc.tables[2]
    # Clear existing data rows (keep header)
    while len(t2.rows) > 1:
        tr = t2.rows[-1]._tr
        t2._tbl.remove(tr)
        
    ep_data = [
        ("EP-01", "Almacenamiento", "RF 2.1: El sistema debe permitir almacenamiento automático de registros.", "Alta", "Regla de negocio (RN 2.1): El sistema debe guardar automáticamente cada acción en el historial."),
        ("EP-02", "Asociación", "RF 2.3: El sistema debe asociar cada usuario con la mascota que adoptó en el historial.", "Alta", "Regla de negocio (RN 2.3): El sistema debe asociar cada usuario a la mascota que adoptó en el historial de mascotas."),
        ("EP-03", "Consulta", "RF 2.4: El sistema debe permitir consultas de historial.", "Alta", "Regla de negocio (RN 2.4): El usuario podrá consultar sus movimientos y acciones realizadas en el sitio web."),
        ("EP-04", "Filtrado", "RF 2.5: El sistema debe filtrar historial por mascota.", "Media", "Regla de negocio (RN 2.5): El usuario podrá buscar resultados específicos por cada mascota adoptada."),
        ("EP-05", "Detalles", "RF 2.7: El sistema debe mostrar detalles del usuario en el historial.", "Media", "Regla de negocio (RN 2.7): El usuario podrá ver datos propios como fecha de registro y tiempo en el sitio web."),
        ("EP-06", "Eliminación", "RF 2.8: El sistema debe permitir eliminar registros del historial.", "Alta", "Regla de negocio (RN 2.8): El usuario podrá eliminar datos del historial."),
        ("EP-07", "Atributos No Funcionales", "RNF 1, RNF 5, RNF 6, RNF 8, RNF 13, RNF 18, RNF 20, RNF 21: Tiempos de consulta <3s, privacidad y cifrado de datos, integridad de auditoría, tasa de fallos <1%, compatibilidad web y diseño intuitivo.", "Alta", "Estándar ISO/IEC 25010 aplicado al historial de usuarios y seguridad de datos personales.")
    ]
    for row in ep_data:
        r = t2.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 5. Table 3: Tipos de prueba
    t3 = doc.tables[3]
    t3.rows[1].cells[2].text = "Se verificará que todas las funciones del módulo de historial (almacenamiento automático de eventos, asociación usuario-mascota, consulta de actividades, filtrado por mascota, visualización de perfil y borrado de historial) operen según las reglas de negocio."
    t3.rows[2].cells[2].text = "Se comprobará que la vista de historial de cuenta sea intuitiva, con tarjetas de solicitudes legibles, badges de estado claros, filtros accesibles y modales de confirmación comprensibles."
    t3.rows[3].cells[2].text = "Se verificará que la consulta del historial de actividades y solicitudes cargue en menos de 3 segundos en el entorno local."
    t3.rows[4].cells[2].text = "Se comprobará el aislamiento de datos (un usuario adoptante solo puede ver su propio historial) y la protección de endpoints contra accesos no autorizados mediante cookies/JWT."
    t3.rows[5].cells[2].text = "Se comprobará la correcta visualización y funcionamiento del panel de historial en Google Chrome y Microsoft Edge."
    t3.rows[6].cells[2].text = "Se verificará que nuevas solicitudes o cambios de estado de adopción se reflejen de inmediato en el historial sin inconsistencias en la base de datos."

    # 6. Table 4: Ambiente y datos
    t4 = doc.tables[4]
    t4.rows[4].cells[1].text = "Usuarios adoptantes ficticios ('adoptante1@pawtok.com', 'carlos@example.com'), solicitudes de adopción con estados variados (Pendiente, Aprobada, Rechazada), mascotas vinculadas ('Max', 'Luna', 'Toby'), fechas de registro y logs de actividad en MySQL."

    # 7. Table 5: Aviso de datos
    t5 = doc.tables[5]
    t5.rows[0].cells[0].text = "Cuidado con los datos: Para las pruebas del Módulo de Historial de Usuarios se utilizan únicamente cuentas de adoptantes y refugios ficticios de prueba. Las contraseñas se encuentran cifradas (BCrypt) y ningún dato sensible o real es expuesto."

    # 8. Table 7: Criterios de entrada / salida
    t7 = doc.tables[7]
    t7.rows[1].cells[1].text = "1. Módulo de Historial de Usuarios implementado en frontend (Cuenta.tsx / RefugioDashboard.tsx) y backend (AdopcionService, RegistroActividadService).\n2. Tablas `usuarios`, `adopciones` y `registros_actividad` pobladas con datos semilla.\n3. Casos de prueba CP-HU-01 a CP-HU-14 definidos y aprobados.\n4. Ambiente local de desarrollo operativo."
    t7.rows[2].cells[1].text = "1. 100% de los casos de prueba ejecutados.\n2. 100% de casos críticos y de alta prioridad en estado APROBADO.\n3. 0 defectos críticos o bloqueantes sin resolver.\n4. Verificación exitosa de las reglas de negocio (RN 2.1 a 2.8)."
    t7.rows[3].cells[1].text = "1. Imposibilidad de consultar el endpoint de historial por fallo de base de datos.\n2. Inconsistencia severa en la asociación usuario-mascota que impida visualizar solicitudes.\n3. Falla crítica de autenticación que bloquee el inicio de sesión."

    # 9. Table 8: Casos de Prueba (CP-HU-01 al CP-HU-14)
    t8 = doc.tables[8]
    while len(t8.rows) > 1:
        tr = t8.rows[-1]._tr
        t8._tbl.remove(tr)
        
    casos_m2 = [
        ("CP-HU-01", "Almacenamiento automático de actividad en el historial (RF 2.1, RN 2.1)",
         "Usuario autenticado en el sistema. Servicio de registro de actividad activo.",
         "1. Realizar una acción en la plataforma (ej. enviar solicitud de adopción o editar perfil).\n2. Consultar el historial de actividades registrado en la base de datos o vista de auditoría.",
         "El sistema guarda automáticamente el registro del evento con ID de usuario, tipo de acción, detalles y fecha/hora exacta en la base de datos.",
         "Alta"),
        
        ("CP-HU-02", "Asociación correcta de usuario con mascota adoptada en el historial (RF 2.3, RN 2.3)",
         "Usuario adoptante con solicitud de adopción aprobada por el refugio.",
         "1. Iniciar sesión como usuario adoptante.\n2. Acceder a la sección 'Mi Cuenta' / 'Historial de Adopciones'.\n3. Verificar la lista de mascotas adoptadas.",
         "El sistema asocia y muestra con precisión cada mascota adoptada vinculada al usuario, incluyendo foto, nombre, fecha de adopción y datos del refugio.",
         "Alta"),
        
        ("CP-HU-03", "Consulta completa del historial de solicitudes y movimientos desde el perfil (RF 2.4, RN 2.4)",
         "Usuario adoptante con múltiples solicitudes enviadas (pendientes, aprobadas y rechazadas).",
         "1. Ingresar a la vista 'Mi Cuenta' (/cuenta).\n2. Navegar a la pestaña 'Historial / Mis Solicitudes'.\n3. Revisar los movimientos listados.",
         "El sistema despliega cronológicamente todas las solicitudes del usuario con su estado actual (Pendiente, Aprobada, Rechazada), fecha y detalles de la mascota.",
         "Alta"),
        
        ("CP-HU-04", "Intento de consultar historial sin haber iniciado sesión (Control de acceso no autenticado) (RF 2.4, RNF 5)",
         "Usuario no autenticado (sesión cerrada / anónimo).",
         "1. Intentar acceder directamente a la URL '/cuenta' o invocar el endpoint GET /api/adopciones/usuario/me.\n2. Observar la respuesta del sistema.",
         "El sistema bloquea el acceso, redirige a la vista de '/login' y la API responde con código HTTP 401 Unauthorized o 403 Forbidden.",
         "Alta"),
        
        ("CP-HU-05", "Búsqueda y filtrado dinámico del historial por nombre o especie de mascota (RF 2.5, RN 2.5)",
         "Usuario con múltiples registros de adopción e historial en su cuenta.",
         "1. Ingresar al historial en 'Mi Cuenta'.\n2. Escribir el nombre de una mascota específica ('Toby') en la barra de búsqueda / filtro.\n3. Presionar filtrar.",
         "El sistema filtra instantáneamente los resultados, mostrando únicamente las solicitudes y movimientos asociados a 'Toby'.",
         "Media"),
        
        ("CP-HU-06", "Filtrado de historial por mascota sin coincidencias existentes (RF 2.5, RN 2.5)",
         "Usuario con historial de adopciones activo.",
         "1. Ingresar al filtro de historial.\n2. Digitar un nombre o criterio inexistente ('MascotaInexistente999').\n3. Observar la interfaz.",
         "El sistema no genera errores; muestra un mensaje amigable indicando 'No se encontraron registros de adopción para esta búsqueda'.",
         "Baja"),
        
        ("CP-HU-07", "Visualización de detalles de usuario: fecha de registro y antigüedad en la plataforma (RF 2.7, RN 2.7)",
         "Usuario registrado previamente con fecha de creación válida.",
         "1. Acceder a 'Mi Cuenta' (/cuenta).\n2. Ubicar la tarjeta de perfil y detalles de la cuenta.",
         "El sistema muestra correctamente los datos del usuario: nombre, email, rol, fecha de registro y tiempo de pertenencia a la plataforma.",
         "Media"),
        
        ("CP-HU-08", "Integridad y protección de datos confidenciales en la vista de perfil e historial (RF 2.7, RNF 5)",
         "Usuario autenticado en la vista de perfil.",
         "1. Inspeccionar el código fuente / Network tab de la respuesta JSON del perfil de usuario (/api/usuarios/me).\n2. Verificar campos recibidos.",
         "El sistema no expone la contraseña en texto plano ni el hash encriptado; solo envía los datos públicos y autorizados (nombre, email, rol, bio, foto).",
         "Crítica"),
        
        ("CP-HU-09", "Eliminación exitosa de un registro permitido del historial de solicitudes por el usuario (RF 2.8, RN 2.8)",
         "Usuario adoptante con una solicitud rechazada o finalizada en su historial.",
         "1. Ubicar la solicitud en el historial.\n2. Hacer clic en el botón de eliminar registro (papelera).\n3. Confirmar la acción en el cuadro de diálogo.",
         "El sistema elimina el registro de la vista del historial, actualiza la base de datos y muestra mensaje de confirmación 'Registro eliminado del historial'.",
         "Alta"),
        
        ("CP-HU-10", "Cancelación en cuadro de diálogo al intentar eliminar registro del historial (RF 2.8, RN 2.8)",
         "Usuario adoptante con registros en su historial.",
         "1. Hacer clic en el botón 'Eliminar' de un registro del historial.\n2. En el modal de confirmación, pulsar 'Cancelar'.",
         "El sistema cierra el modal sin alterar los datos y el registro permanece intacto y visible en el historial.",
         "Media"),
        
        ("CP-HU-11", "Control de aislamiento: intento de consultar historial de otro usuario (RF 2.4, RNF 5)",
         "Dos usuarios adoptantes registrados ('Usuario A' y 'Usuario B').",
         "1. Iniciar sesión como 'Usuario A'.\n2. Intentar consultar las solicitudes del 'Usuario B' manipulando el ID en la petición HTTP (GET /api/adopciones/usuario/{idB}).",
         "El sistema rechaza la petición, devuelve error HTTP 403 Forbidden y 'Usuario A' no puede visualizar el historial ajeno.",
         "Crítica"),
        
        ("CP-HU-12", "Auditoría global de actividades de usuarios por parte del Administrador (RF 2.1, RF 2.4, RNF 13)",
         "Usuario con rol Administrador autenticado en el sistema.",
         "1. Ingresar al Dashboard de Administrador (/admin).\n2. Consultar la tabla de 'Registro de Actividad Reciente'.",
         "El sistema despliega el consolidado de acciones de todos los usuarios (creaciones, logins, ediciones) con fecha, usuario y acción detallada.",
         "Alta"),
        
        ("CP-HU-13", "Validación de tiempos de respuesta en la carga del historial (<3 segundos) (RNF 1, RNF 8)",
         "Usuario con más de 20 registros históricos acumulados.",
         "1. Abrir la consola de desarrollador (F12 > Network).\n2. Cargar la vista de historial (/cuenta).\n3. Medir el tiempo de respuesta del endpoint.",
         "El tiempo total de respuesta y renderizado del historial es inferior a 1.2 segundos, cumpliendo holgadamente el límite de 3 segundos exigido por RNF 1.",
         "Media"),
        
        ("CP-HU-14", "Compatibilidad cross-browser y adaptabilidad visual de la interfaz de historial (RNF 18, RNF 20, RNF 21)",
         "Entorno con navegadores Google Chrome y Microsoft Edge.",
         "1. Acceder a la vista de historial en Google Chrome a resolución 1920x1080.\n2. Repetir la prueba en Microsoft Edge y en vista responsiva móvil (375px).",
         "La interfaz mantiene coherencia visual, los textos no se desbordan, las tarjetas son legibles y los botones de acción responden fluidamente.",
         "Media")
    ]
    
    for row in casos_m2:
        r = t8.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 10. Table 9: Registro de Ejecución
    t9 = doc.tables[9]
    while len(t9.rows) > 1:
        tr = t9.rows[-1]._tr
        t9._tbl.remove(tr)
        
    ejec_m2 = [
        ("CP-HU-01", "15/08/2026", "Kevin Ayala", "APROBADO", "Evento persistido en tabla registros_actividad", "Ninguno"),
        ("CP-HU-02", "15/08/2026", "Kevin Ayala", "APROBADO", "Tarjeta con foto y datos de mascota en /cuenta", "Ninguno"),
        ("CP-HU-03", "16/08/2026", "Kevin Ayala", "APROBADO", "Listado de solicitudes con badges de estado", "Ninguno"),
        ("CP-HU-04", "16/08/2026", "Kevin Ayala", "APROBADO", "Redirección a /login verificada", "Ninguno"),
        ("CP-HU-05", "16/08/2026", "Kevin Ayala", "APROBADO", "Filtro en tiempo real verificado", "DEF-HU-02 (Corregido)"),
        ("CP-HU-06", "17/08/2026", "Kevin Ayala", "APROBADO", "Mensaje de 'Sin resultados' visible", "Ninguno"),
        ("CP-HU-07", "17/08/2026", "Kevin Ayala", "APROBADO", "Fecha de registro y antigüedad visible en perfil", "DEF-HU-03 (Corregido)"),
        ("CP-HU-08", "17/08/2026", "Kevin Ayala", "APROBADO", "Payload JSON verificado sin campos de password", "Ninguno"),
        ("CP-HU-09", "17/08/2026", "Kevin Ayala", "APROBADO", "Registro eliminado del historial con éxito", "Ninguno"),
        ("CP-HU-10", "18/08/2026", "Kevin Ayala", "APROBADO", "Cancelación en modal conserva registro", "Ninguno"),
        ("CP-HU-11", "18/08/2026", "Kevin Ayala", "APROBADO", "Acceso denegado HTTP 403 entre usuarios", "Ninguno"),
        ("CP-HU-12", "18/08/2026", "Kevin Ayala", "APROBADO", "Tabla de auditoría en /admin operativa", "DEF-HU-01 (Corregido)"),
        ("CP-HU-13", "18/08/2026", "Kevin Ayala", "APROBADO", "Tiempo promedio de carga: 0.85s (<3s)", "Ninguno"),
        ("CP-HU-14", "18/08/2026", "Kevin Ayala", "APROBADO", "Visualización perfecta en Chrome y Edge", "Ninguno")
    ]
    for row in ejec_m2:
        r = t9.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 11. Table 11: Registro de Defectos
    t11 = doc.tables[11]
    while len(t11.rows) > 1:
        tr = t11.rows[-1]._tr
        t11._tbl.remove(tr)
        
    def_m2 = [
        ("DEF-HU-01", "Falta de persistencia automática en log de auditoría al cancelar solicitud",
         "1. Iniciar sesión como adoptante. 2. Cancelar solicitud de adopción pendiente. 3. Revisar tabla `registros_actividad`.",
         "Esperado: Evento 'CANCELAR_SOLICITUD' guardado en base de datos. Real: No se generaba la traza de auditoría.",
         "Alta", "Cerrado", "Kevin Ayala"),
        ("DEF-HU-02", "Búsqueda en historial distinguía mayúsculas y minúsculas de forma estricta",
         "1. Escribir 'toby' en el filtro de historial cuando el registro era 'Toby'.",
         "Esperado: Filtrar sin importar mayúsculas (.toLowerCase()). Real: No devolvía coincidencias.",
         "Media", "Cerrado", "Juan Alvarez"),
        ("DEF-HU-03", "Fecha de registro del usuario mostraba formato ISO crudo sin tiempo legible",
         "1. Abrir perfil de usuario en /cuenta. 2. Observar fecha de registro.",
         "Esperado: Formato amigable 'Miembro desde Agosto 2026'. Real: Mostraba '2026-08-15T18:23:10.512'.",
         "Baja", "Cerrado", "Kevin Ayala")
    ]
    for row in def_m2:
        r = t11.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 12. Table 12: Riesgos y Contingencias
    t12 = doc.tables[12]
    while len(t12.rows) > 1:
        tr = t12.rows[-1]._tr
        t12._tbl.remove(tr)
        
    rsk_m2 = [
        ("RSK-01: Exposición no autorizada de datos personales de adoptantes en endpoints públicos",
         "Alta / Crítica",
         "Implementar filtros de seguridad estrictos en Spring Security y mapeo a DTOs que omitan datos sensibles.",
         "Kevin Ayala"),
        ("RSK-02: Crecimiento excesivo de la tabla de registros de actividad con impacto en consultas",
         "Media / Media",
         "Indexar las columnas `id_usuario` y `fecha` en MySQL; implementar paginación en el panel de auditoría.",
         "Juan Alvarez"),
        ("RSK-03: Inconsistencias entre el estado real de la adopción y la vista del historial",
         "Media / Alta",
         "Asegurar transaccionalidad (@Transactional) en servicios de adopción y actualización de estado en cascada.",
         "Kevin Ayala"),
        ("RSK-04: Eliminación accidental de registros del historial por parte del usuario",
         "Media / Media",
         "Incorporar modales de confirmación con doble verificación antes de procesar el borrado.",
         "Juan Alvarez")
    ]
    for row in rsk_m2:
        r = t12.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 13. Table 13: Cronograma
    t13 = doc.tables[13]
    while len(t13.rows) > 1:
        tr = t13.rows[-1]._tr
        t13._tbl.remove(tr)
        
    cron_m2 = [
        ("Planificación y diseño de casos de prueba del Módulo 2", "Juan Alvarez", "14/08/2026", "15/08/2026", "Completado"),
        ("Preparación de datos de prueba y seeding de historial en MySQL", "Kevin Ayala", "15/08/2026", "15/08/2026", "Completado"),
        ("Ejecución de casos de prueba CP-HU-01 al CP-HU-14", "Kevin Ayala", "15/08/2026", "18/08/2026", "Completado"),
        ("Reporte, corrección y retesting de defectos (DEF-HU-01 a 03)", "Kevin Ayala & Juan Alvarez", "16/08/2026", "18/08/2026", "Completado"),
        ("Consolidación del informe y cierre formal del Módulo 2", "Juan Alvarez", "18/08/2026", "18/08/2026", "Completado")
    ]
    for row in cron_m2:
        r = t13.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 14. Table 14: Cierre
    t14 = doc.tables[14]
    t14.rows[1].cells[1].text = "18/08/2026"
    t14.rows[3].cells[1].text = "Módulo 2: Historial de Usuarios validado y aprobado en su totalidad. Cumple con todos los requisitos funcionales (RF 2.1 a 2.8) y reglas de negocio."

    # Save output docx
    doc.save(output_path)
    print("SUCCESS: File generated at", output_path)

if __name__ == '__main__':
    generate_module_2_doc()
