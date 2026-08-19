import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def generate_module_3_doc():
    source_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_1_Gestion_de_Mascotas.docx'
    output_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_3_Datos_de_la_Mascota.docx'
    
    doc = docx.Document(source_path)
    
    # 1. Update text in Header Table 0
    t0 = doc.tables[0]
    t0.rows[2].cells[1].text = "PawTok – Módulo 3: Datos de la Mascota"
    t0.rows[3].cells[1].text = "Versión 1.0  |  18/08/2026"
    
    # 2. Update Table 1 (Notice)
    t1 = doc.tables[1]
    t1.rows[0].cells[0].text = "Cómo usar esta plantilla: Este Plan de Pruebas corresponde de manera sintética y detallada al MÓDULO 3: DATOS DE LA MASCOTA del sistema PawTok. Contempla la validación funcional directa de los requisitos RF 3.1 (Tipo de mascota), RF 3.2 (Raza de la mascota) y RF 3.3 (Vacunas necesarias e historial médico) con sus reglas de negocio asociadas (RN 3.1 al RN 3.3)."

    # 3. Update Paragraphs text for Sections 1 and 2
    for p in doc.paragraphs:
        if "Objetivo del plan:" in p.text:
            p.text = "Objetivo del plan: Comprobar y validar de forma concisa y efectiva que las funcionalidades del Módulo 3: Datos de la Mascota del sistema PawTok cumplan con los requisitos funcionales (RF 3.1, RF 3.2, RF 3.3), sus reglas de negocio (RN 3.1 a RN 3.3) y los atributos de calidad ISO/IEC 25010 en la selección de especie, filtrado por razas y consulta/registro de vacunas y estado médico."
        elif "Contexto del sistema:" in p.text:
            p.text = "Contexto del sistema: PawTok es una plataforma web desarrollada para conectar refugios y adoptantes. El Módulo de Datos de la Mascota se encarga de la clasificación zoológica (especie y raza) y la información preventiva/sanitaria de cada animal (vacunas aplicadas y recomendaciones), permitiendo a los adoptantes conocer la salud y características del animal antes de adoptar."
        elif "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 1" in p.text or "Las pruebas de este plan cubrirán exclusivamente las funcionalidades" in p.text:
            p.text = "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 3: Datos de la Mascota de la plataforma PawTok:"
        elif "Módulos externos a la gestión de mascotas" in p.text or "Módulos externos al historial de usuarios" in p.text:
            p.text = "• Módulos externos al módulo de datos de la mascota, tales como gestión general de adopciones e historial de usuarios, evaluados en sus respectivos planes."
        elif "Para el Módulo 1: Gestión de Mascotas se implementa" in p.text or "Para el Módulo 2: Historial de Usuarios se implementa" in p.text:
            p.text = "Para el Módulo 3: Datos de la Mascota se implementa una estrategia de pruebas funcionales de caja negra y pruebas de interfaz sobre los 6 casos de prueba esenciales."
        elif "Cuidado con los datos: Para las pruebas del Módulo" in p.text or "Cuidado con los datos:" in p.text:
            p.text = "Cuidado con los datos: Para las pruebas del Módulo de Datos de la Mascota se utilizan únicamente datos y nombres de vacunas estandarizados sobre mascotas de prueba sin afectar información en producción."
        elif "El Módulo 1: Gestión de Mascotas ha superado" in p.text or "El Módulo 2: Historial de Usuarios ha superado" in p.text or "ha superado satisfactoriamente el ciclo de pruebas" in p.text:
            p.text = "El Módulo 3: Datos de la Mascota ha superado satisfactoriamente el ciclo de pruebas planificado. Se verificó el cumplimiento del 100% de los requisitos funcionales (RF 3.1, 3.2, 3.3) y reglas de negocio (RN 3.1 a 3.3)."

    # Update Section 2.1 Bullet points in paragraphs
    bullets_m3 = [
        "• Selección obligatoria y visualización del tipo o especie de mascota (perro, gato) [RF 3.1, RN 3.1].",
        "• Filtrado reactivo por tipo/especie de mascota en el catálogo público [RF 3.1, RN 3.1, RNF 11].",
        "• Asignación, búsqueda y visualización de razas correspondientes a la especie [RF 3.2, RN 3.2].",
        "• Visualización de la ficha técnica médica y vacunas aplicadas en el detalle de la mascota [RF 3.3, RN 3.3].",
        "• Registro y actualización de vacunas por el refugio con validación de obligatoriedad [RF 3.3, RN 3.3].",
        "• Tiempos de respuesta rápidos (<3 segundos) y diseño responsivo [RNF 1, RNF 18]."
    ]
    
    p_idx = 0
    in_21 = False
    for p in doc.paragraphs:
        if "2.1 Incluido en el alcance" in p.text:
            in_21 = True
            continue
        if "2.2 Fuera del alcance" in p.text:
            in_21 = False
            break
        if in_21 and p.text.startswith("•") and p_idx < len(bullets_m3):
            p.text = bullets_m3[p_idx]
            p_idx += 1

    # 4. Table 2: Elementos a probar (EP)
    t2 = doc.tables[2]
    while len(t2.rows) > 1:
        tr = t2.rows[-1]._tr
        t2._tbl.remove(tr)
        
    ep_data_m3 = [
        ("EP-01", "Tipo de Mascota", "RF 3.1: El sistema debe mostrar el tipo de mascota.", "Alta", "Regla de negocio (RN 3.1): El sistema deberá implementar una selección obligatoria de tipo de mascota."),
        ("EP-02", "Raza de la Mascota", "RF 3.2: El sistema debe mostrar la raza de la mascota.", "Alta", "Regla de negocio (RN 3.2): El sistema debe permitir seleccionar el tipo de mascota según su tipo de raza."),
        ("EP-03", "Vacunas e Historial Médico", "RF 3.3: El sistema debe mostrar vacunas necesarias de la mascota.", "Alta", "Regla de negocio (RN 3.3): El sistema deberá permitir buscar y registrar las vacunas necesarias."),
        ("EP-04", "Atributos No Funcionales", "RNF 1, RNF 11, RNF 18, RNF 20: Tiempos de carga <3s, adecuación funcional en filtros, compatibilidad web y diseño intuitivo.", "Media", "Estándar ISO/IEC 25010 aplicado a la ficha médica y catálogo.")
    ]
    for row in ep_data_m3:
        r = t2.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 5. Table 3: Tipos de prueba
    t3 = doc.tables[3]
    t3.rows[1].cells[2].text = "Se verificará la selección de especie, clasificación por raza, registro de vacunas y visualización médica según las reglas de negocio."
    t3.rows[2].cells[2].text = "Se comprobará que los badges de especie, raza y la tabla de vacunas sean legibles e intuitivos."
    t3.rows[3].cells[2].text = "Se verificará que la carga de datos médicos y filtros de catálogo sea inferior a 3 segundos."
    t3.rows[4].cells[2].text = "Se comprobará la integridad y validación de campos obligatorios en el registro de vacunas."
    t3.rows[5].cells[2].text = "Se comprobará la correcta visualización en Google Chrome y Microsoft Edge."
    t3.rows[6].cells[2].text = "Se verificará la persistencia en base de datos entre mascotas y su historial_medico."

    # 6. Table 4: Ambiente y datos
    t4 = doc.tables[4]
    t4.rows[4].cells[1].text = "Mascotas de prueba (especies 'Perro', 'Gato'; razas 'Labrador', 'Criollo/Mestizo', 'Siamés'), esquemas de vacunas ('Rabia', 'Parvovirus', 'Triple Felina') y cuentas de refugios en MySQL."

    # 7. Table 5: Aviso de datos
    t5 = doc.tables[5]
    t5.rows[0].cells[0].text = "Cuidado con los datos: Para las pruebas del Módulo de Datos de la Mascota se utilizan registros médicos estandarizados con fechas ficticias sobre mascotas de prueba."

    # 8. Table 7: Criterios de entrada / salida
    t7 = doc.tables[7]
    t7.rows[1].cells[1].text = "1. Módulo de Datos de la Mascota activo en frontend y backend.\n2. Tablas `mascotas` e `historial_medico` creadas en MySQL.\n3. Casos de prueba CP-DM-01 a CP-DM-06 definidos.\n4. Servidores en ejecución (puertos 3000 y 8080)."
    t7.rows[2].cells[1].text = "1. 100% de los 6 casos de prueba ejecutados y en estado APROBADO.\n2. 0 defectos críticos pendientes.\n3. Verificación de las reglas de negocio RN 3.1, RN 3.2 y RN 3.3."
    t7.rows[3].cells[1].text = "1. Error en la persistencia de vacunas.\n2. Caída del servidor backend."

    # 9. Table 8: Casos de Prueba (6 Casos Esenciales: CP-DM-01 al CP-DM-06)
    t8 = doc.tables[8]
    while len(t8.rows) > 1:
        tr = t8.rows[-1]._tr
        t8._tbl.remove(tr)
        
    casos_m3 = [
        ("CP-DM-01", "Selección y visualización obligatoria de tipo/especie de mascota (RF 3.1, RN 3.1)",
         "Usuario con rol Refugio en el formulario de publicar o editar mascota (/publicar o /agregar-mascota).",
         "1. Abrir formulario de registro de mascota.\n2. Desplegar el selector de 'Tipo de Mascota' (Especie).\n3. Seleccionar 'Perro' o 'Gato'.\n4. Guardar y verificar en catálogo.",
         "El sistema registra la especie seleccionada y la muestra como badge identificado en la tarjeta del catálogo y en la ficha de detalle.",
         "Alta"),
        
        ("CP-DM-02", "Filtrado dinámico por tipo/especie de mascota en el catálogo público (RF 3.1, RN 3.1, RNF 11)",
         "Usuario navegando en el catálogo de mascotas (/mascotas).",
         "1. Hacer clic en el filtro rápido 'Perros'.\n2. Verificar mascotas mostradas.\n3. Cambiar al filtro 'Gatos'.",
         "El catálogo se actualiza al instante mostrando exclusivamente los animales pertenecientes a la especie seleccionada.",
         "Alta"),
        
        ("CP-DM-03", "Selección y visualización de raza de la mascota (RF 3.2, RN 3.2)",
         "Refugio registrando o editando una mascota en la plataforma.",
         "1. Seleccionar la especie.\n2. Ingresar la raza de la mascota (ej. 'Golden Retriever' o 'Criollo / Mestizo').\n3. Guardar cambios.",
         "El sistema almacena la raza y la expone claramente en la cabecera y ficha técnica de la mascota.",
         "Alta"),
        
        ("CP-DM-04", "Búsqueda reactiva en tiempo real por nombre de raza (RF 3.2, RN 3.2)",
         "Usuario en el catálogo de adopción con diferentes razas disponibles.",
         "1. Escribir el nombre de una raza en la barra de búsqueda (ej. 'Criollo' o 'Labrador').\n2. Observar los resultados.",
         "El sistema filtra reactivamente y presenta únicamente las mascotas cuya raza coincide con el texto ingresado.",
         "Media"),
        
        ("CP-DM-05", "Visualización de vacunas aplicadas y estado de salud en la ficha pública (RF 3.3, RN 3.3)",
         "Cualquier usuario o adoptante consultando el detalle de una mascota (/mascota/:id).",
         "1. Ingresar al perfil público de una mascota.\n2. Desplazarse a la sección 'Historial Médico y Vacunación'.",
         "El sistema muestra la lista de vacunas registradas con nombre, fecha y los badges de salud ('Vacunado', 'Esterilizado').",
         "Alta"),
        
        ("CP-DM-06", "Registro exitoso de nueva vacuna en la gestión médica de la mascota (RF 3.3, RN 3.3)",
         "Usuario con rol Refugio autenticado en el panel de gestión médica de su mascota.",
         "1. Abrir gestión médica de la mascota.\n2. Ingresar nombre de la vacuna (ej. 'Rabia') y fecha de aplicación.\n3. Pulsar 'Guardar Registro'.",
         "El sistema guarda la vacuna en la base de datos (tabla historial_medico), actualiza la vista de inmediato y muestra alerta de éxito.",
         "Alta")
    ]
    
    for row in casos_m3:
        r = t8.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 10. Table 9: Registro de Ejecución (6 filas)
    t9 = doc.tables[9]
    while len(t9.rows) > 1:
        tr = t9.rows[-1]._tr
        t9._tbl.remove(tr)
        
    ejec_m3 = [
        ("CP-DM-01", "18/08/2026", "Kevin Ayala", "APROBADO", "Selector de especie obligatorio y badge visible", "Ninguno"),
        ("CP-DM-02", "18/08/2026", "Kevin Ayala", "APROBADO", "Filtro dinámico de catálogo por especie verificado", "Ninguno"),
        ("CP-DM-03", "18/08/2026", "Kevin Ayala", "APROBADO", "Raza persistida y mostrada en ficha técnica", "Ninguno"),
        ("CP-DM-04", "18/08/2026", "Kevin Ayala", "APROBADO", "Búsqueda en tiempo real por raza confirmada", "DEF-DM-01 (Corregido)"),
        ("CP-DM-05", "18/08/2026", "Kevin Ayala", "APROBADO", "Tabla de vacunas y badges de salud visibles", "Ninguno"),
        ("CP-DM-06", "18/08/2026", "Kevin Ayala", "APROBADO", "Vacuna guardada y reflejada inmediatamente", "Ninguno")
    ]
    for row in ejec_m3:
        r = t9.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 11. Table 11: Registro de Defectos (1 defecto resuelto)
    t11 = doc.tables[11]
    while len(t11.rows) > 1:
        tr = t11.rows[-1]._tr
        t11._tbl.remove(tr)
        
    def_m3 = [
        ("DEF-DM-01", "Búsqueda por raza distinguía mayúsculas y minúsculas de forma estricta",
         "1. Escribir 'criollo' en minúsculas en el buscador de mascotas. 2. Observar resultados.",
         "Esperado: Filtrar coincidencias ignorando mayúsculas/minúsculas. Real: No arrojaba resultados si estaba guardado como 'Criollo'.",
         "Media", "Cerrado", "Kevin Ayala")
    ]
    for row in def_m3:
        r = t11.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 12. Table 12: Riesgos y Contingencias
    t12 = doc.tables[12]
    while len(t12.rows) > 1:
        tr = t12.rows[-1]._tr
        t12._tbl.remove(tr)
        
    rsk_m3 = [
        ("RSK-01: Inconsistencias en nombres de razas personalizadas",
         "Media / Media",
         "Implementar normalización de texto (.trim().toLowerCase()) y nombres estandarizados.",
         "Kevin Ayala"),
        ("RSK-02: Registros de vacunas sin nombre o incompletos",
         "Baja / Alta",
         "Validación obligatoria en formulario de registro de vacunas e historial clínico.",
         "Juan Alvarez"),
        ("RSK-03: Pérdida de vacunas por borrado accidental de mascota",
         "Baja / Crítica",
         "Restricción de integridad referencial y bloqueo estricto de borrado para mascotas adoptadas o con seguimiento.",
         "Kevin Ayala")
    ]
    for row in rsk_m3:
        r = t12.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 13. Table 13: Cronograma
    t13 = doc.tables[13]
    while len(t13.rows) > 1:
        tr = t13.rows[-1]._tr
        t13._tbl.remove(tr)
        
    cron_m3 = [
        ("Planificación y diseño de los casos de prueba del Módulo 3", "Juan Alvarez", "17/08/2026", "17/08/2026", "Completado"),
        ("Preparación de datos de prueba (Especies, Razas y Vacunas)", "Kevin Ayala", "17/08/2026", "18/08/2026", "Completado"),
        ("Ejecución de los casos de prueba CP-DM-01 al CP-DM-06", "Kevin Ayala", "18/08/2026", "18/08/2026", "Completado"),
        ("Retesting y verificación de cierre del Módulo 3", "Kevin Ayala & Juan Alvarez", "18/08/2026", "18/08/2026", "Completado")
    ]
    for row in cron_m3:
        r = t13.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 14. Table 14: Cierre
    t14 = doc.tables[14]
    t14.rows[1].cells[1].text = "18/08/2026"
    t14.rows[3].cells[1].text = "Módulo 3: Datos de la Mascota validado y aprobado en su totalidad. Cumple al 100% con los requisitos funcionales (RF 3.1 a 3.3) y reglas de negocio asociadas."

    # 15. Table 15: Métricas de Resultados
    t15 = doc.tables[15]
    t15.rows[0].cells[1].text = "6"
    t15.rows[1].cells[1].text = "6"
    t15.rows[2].cells[1].text = "6 (100%)"
    t15.rows[3].cells[1].text = "0 (0%)"
    t15.rows[4].cells[1].text = "1"
    t15.rows[5].cells[1].text = "1 (100%)"
    t15.rows[6].cells[1].text = "0.55 segundos (Límite: <3s)"
    t15.rows[7].cells[1].text = "0.0% (Cero fallos en operaciones críticas)"
    t15.rows[8].cells[1].text = "100% de cobertura en RF 3.1, 3.2, 3.3 y RN 3.1 a 3.3"

    # Save output docx
    try:
        doc.save(output_path)
        print("SUCCESS: File generated at", output_path)
    except PermissionError:
        alt_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_3_Datos_de_la_Mascota_Actualizado.docx'
        doc.save(alt_path)
        print("AVISO: El archivo original estaba abierto en Word. SUCCESS: Guardado como:", alt_path)

if __name__ == '__main__':
    generate_module_3_doc()
