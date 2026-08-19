import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def generate_module_4_doc():
    source_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_1_Gestion_de_Mascotas.docx'
    output_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_4_Administracion_del_Sistema.docx'
    
    doc = docx.Document(source_path)
    
    # 1. Update text in Header Table 0
    t0 = doc.tables[0]
    t0.rows[2].cells[1].text = "PawTok – Módulo 4: Administración del Sistema"
    t0.rows[3].cells[1].text = "Versión 1.0  |  18/08/2026"
    
    # 2. Update Table 1 (Notice)
    t1 = doc.tables[1]
    t1.rows[0].cells[0].text = "Cómo usar esta plantilla: Este Plan de Pruebas corresponde de manera exclusiva y detallada al MÓDULO 4: ADMINISTRACIÓN DEL SISTEMA Y SUGERENCIAS DE CUIDADO del sistema PawTok. Contempla la validación funcional de los requisitos RF 4.1 (Sugerencias de cuidado), RF 4.2 (Calificación del cuidado), RF 4.4 (Listado de usuarios y refugios) y RF 4.5 (Activación, aprobación y suspensión de cuentas), junto con las reglas de negocio asociadas (RN 4.1 al RN 4.5)."

    # 3. Update Paragraphs text for Sections 1 and 2
    for p in doc.paragraphs:
        if "Objetivo del plan:" in p.text:
            p.text = "Objetivo del plan: Comprobar y validar de manera rigurosa que las funcionalidades del Módulo 4: Administración del Sistema de PawTok cumplan con los requisitos funcionales (RF 4.1, RF 4.2, RF 4.4, RF 4.5), sus reglas de negocio (RN 4.1 a RN 4.5) y los estándares ISO/IEC 25010 en la gestión global de usuarios, verificación de refugios aliados, auditoría de actividades y retroalimentación de cuidados post-adopción."
        elif "Contexto del sistema:" in p.text:
            p.text = "Contexto del sistema: PawTok es una plataforma web desarrollada para conectar refugios y adoptantes. El Módulo de Administración del Sistema es el centro de control que permite a los usuarios con rol Administrador supervisar métricas globales, gestionar cuentas de adoptantes y organizaciones aliadas, aprobar o rechazar solicitudes de refugio, monitorear la auditoría de operaciones y brindar sugerencias y evaluaciones de cuidado en los seguimientos."
        elif "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 1" in p.text or "Las pruebas de este plan cubrirán exclusivamente las funcionalidades" in p.text:
            p.text = "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 4: Administración del Sistema de la plataforma PawTok:"
        elif "Módulos externos a la gestión de mascotas" in p.text or "Módulos externos al historial de usuarios" in p.text:
            p.text = "• Módulos externos a la administración del sistema, tales como la publicación básica de mascotas (Módulo 1), historial particular de adoptantes (Módulo 2) y catálogo de razas (Módulo 3), evaluados en sus planes específicos."
        elif "Para el Módulo 1: Gestión de Mascotas se implementa" in p.text or "Para el Módulo 2: Historial de Usuarios se implementa" in p.text:
            p.text = "Para el Módulo 4: Administración del Sistema se implementa una estrategia de pruebas funcionales de caja negra y pruebas de seguridad de roles para verificar el control de acceso, la gestión de usuarios y la persistencia de sugerencias y calificaciones."
        elif "Cuidado con los datos: Para las pruebas del Módulo" in p.text or "Cuidado con los datos:" in p.text:
            p.text = "Cuidado con los datos: Para las pruebas del Módulo de Administración se utilizan cuentas de prueba con roles ADMIN, REFUGIO y USUARIO, garantizando que ninguna cuenta real o contraseña sea vulnerada."
        elif "El Módulo 1: Gestión de Mascotas ha superado" in p.text or "El Módulo 2: Historial de Usuarios ha superado" in p.text or "ha superado satisfactoriamente el ciclo de pruebas" in p.text:
            p.text = "El Módulo 4: Administración del Sistema ha superado satisfactoriamente el ciclo de pruebas planificado. Se verificó el cumplimiento del 100% de los requisitos funcionales (RF 4.1, 4.2, 4.4, 4.5) y reglas de negocio (RN 4.1 a 4.5)."

    # Update Section 2.1 Bullet points in paragraphs
    bullets_m4 = [
        "• Visualización de métricas globales y listado completo de usuarios y refugios registrados [RF 4.4, RN 4.4].",
        "• Activación, suspensión y eliminación segura de cuentas de usuario desde el panel de control [RF 4.5, RN 4.5].",
        "• Verificación, aprobación y rechazo de solicitudes de registro de nuevos refugios [RF 4.5, RN 4.5].",
        "• Emisión de sugerencias y pautas de cuidado hacia los adoptantes en el seguimiento post-adopción [RF 4.1, RN 4.1].",
        "• Calificación y puntuación del nivel de bienestar y cuidado de la mascota adoptada [RF 4.2, RN 4.2].",
        "• Control de seguridad y protección de rutas administrativas (HTTP 403 Forbidden a adoptantes) [RNF 5].",
        "• Tiempos de carga de métricas y tablas inferiores a 3 segundos [RNF 1, RNF 8]."
    ]
    
    p_idx = 0
    in_21 = False
    for p in doc.paragraphs:
        if "2.1 Incluido en el alcance" in p.text:
            in_21 = True
            continue
        if "2.2 Fuera del alcance" in p.text:
            in_21 = False
            break
        if in_21 and p.text.startswith("•") and p_idx < len(bullets_m4):
            p.text = bullets_m4[p_idx]
            p_idx += 1

    # 4. Table 2: Elementos a probar (EP)
    t2 = doc.tables[2]
    while len(t2.rows) > 1:
        tr = t2.rows[-1]._tr
        t2._tbl.remove(tr)
        
    ep_data_m4 = [
        ("EP-01", "Listado y Métricas de Usuarios", "RF 4.4: El sistema debe mostrar la lista de usuarios y métricas globales.", "Alta", "Regla de negocio (RN 4.4): El administrador podrá visualizar todas las cuentas activas."),
        ("EP-02", "Gestión y Suspensión de Cuentas", "RF 4.5: El sistema debe activar o desactivar cuentas de usuarios y refugios.", "Alta", "Regla de negocio (RN 4.5): El administrador podrá suspender o eliminar cuentas si es necesario."),
        ("EP-03", "Sugerencias de Cuidado", "RF 4.1: El sistema debe dar sugerencias de cuidados para mascotas.", "Media", "Regla de negocio (RN 4.1): El administrador/refugio podrá dar sugerencias al usuario para el cuidado."),
        ("EP-04", "Calificación de Bienestar", "RF 4.2: El sistema debe calificar el cuidado de la mascota.", "Media", "Regla de negocio (RN 4.2): El administrador/refugio podrá puntuar al usuario según el nivel de cuidado."),
        ("EP-05", "Atributos No Funcionales", "RNF 1, RNF 5, RNF 8, RNF 13, RNF 18: Tiempos de carga <3s, seguridad de roles, auditoría de operaciones y compatibilidad web.", "Alta", "Estándar ISO/IEC 25010 aplicado al panel de administración.")
    ]
    for row in ep_data_m4:
        r = t2.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 5. Table 3: Tipos de prueba
    t3 = doc.tables[3]
    t3.rows[1].cells[2].text = "Se verificará la visualización de métricas, gestión de cuentas de usuario, verificación de refugios y emisión de sugerencias de cuidado según las reglas de negocio."
    t3.rows[2].cells[2].text = "Se comprobará que el panel de administración sea intuitivo, con tarjetas de estadísticas legibles, tablas ordenadas y botones de acción claros."
    t3.rows[3].cells[2].text = "Se verificará que las estadísticas y listados de administración carguen en menos de 3 segundos."
    t3.rows[4].cells[2].text = "Se comprobará el bloqueo estricto de rutas /admin a usuarios no autenticados o con rol USUARIO."
    t3.rows[5].cells[2].text = "Se comprobará la visualización fluida del panel administrativo en Google Chrome y Microsoft Edge."
    t3.rows[6].cells[2].text = "Se verificará la persistencia inmediata en base de datos al suspender cuentas o aprobar solicitudes de refugio."

    # 6. Table 4: Ambiente y datos
    t4 = doc.tables[4]
    t4.rows[4].cells[1].text = "Cuenta de Administrador ('admin@pawtok.com'), cuentas de refugios pendientes y aprobados, usuarios adoptantes registrados y registros de seguimiento en MySQL."

    # 7. Table 5: Aviso de datos
    t5 = doc.tables[5]
    t5.rows[0].cells[0].text = "Cuidado con los datos: Para las pruebas del Módulo de Administración se utilizan cuentas de prueba con roles asignados sin comprometer datos confidenciales reales."

    # 8. Table 7: Criterios de entrada / salida
    t7 = doc.tables[7]
    t7.rows[1].cells[1].text = "1. Módulo de Administración implementado en frontend (AdminDashboard.tsx, AdminSolicitudesRefugio.tsx) y backend (AdminController).\n2. Tablas `usuarios`, `refugios` y `seguimientos` pobladas con datos de prueba.\n3. Casos de prueba CP-AS-01 a CP-AS-06 aprobados.\n4. Servidores en ejecución (puertos 3000 y 8080)."
    t7.rows[2].cells[1].text = "1. 100% de los 6 casos de prueba ejecutados y en estado APROBADO.\n2. 0 defectos críticos pendientes.\n3. Verificación de las reglas de negocio RN 4.1, RN 4.2, RN 4.4 y RN 4.5."
    t7.rows[3].cells[1].text = "1. Falla de seguridad que permita a adoptantes ingresar al panel de admin.\n2. Inconsistencia en la aprobación de refugios."

    # 9. Table 8: Casos de Prueba (6 Casos Esenciales: CP-AS-01 al CP-AS-06)
    t8 = doc.tables[8]
    while len(t8.rows) > 1:
        tr = t8.rows[-1]._tr
        t8._tbl.remove(tr)
        
    casos_m4 = [
        ("CP-AS-01", "Visualización global de estadísticas y listado de usuarios en el Dashboard de Admin (RF 4.4, RN 4.4)",
         "Usuario con rol Administrador autenticado en la plataforma.",
         "1. Ingresar a la ruta '/admin'.\n2. Observar las tarjetas de estadísticas globales (Total Usuarios, Adoptantes, Refugios, Mascotas).\n3. Consultar la tabla de usuarios registrados.",
         "El sistema despliega en tiempo real las cifras consolidadas de la plataforma y lista todos los usuarios con su nombre, correo, rol y fecha de creación.",
         "Alta"),
        
        ("CP-AS-02", "Gestión y eliminación segura de cuentas de usuario por el Administrador (RF 4.5, RN 4.5)",
         "Administrador navegando en la tabla de usuarios en /admin.",
         "1. Ubicar un usuario de prueba en la tabla.\n2. Hacer clic en el botón de eliminar/suspender usuario.\n3. Confirmar la acción en el modal.",
         "El sistema elimina/suspende la cuenta del usuario en la base de datos, actualiza la tabla de inmediato y muestra alerta de éxito.",
         "Alta"),
        
        ("CP-AS-03", "Verificación y aprobación de solicitudes de registro de refugios aliados (RF 4.5, RN 4.5)",
         "Refugio recién registrado en estado 'Pendiente' y Administrador en /admin/solicitudes-refugio.",
         "1. Acceder al panel de solicitudes de refugio.\n2. Revisar la información y documentos del refugio.\n3. Pulsar 'Aprobar Refugio'.",
         "El sistema actualiza el estado del refugio a 'Aprobado', asigna el rol REFUGIO al usuario y habilita sus permisos para publicar mascotas.",
         "Alta"),
        
        ("CP-AS-04", "Registro y consulta de sugerencias de cuidado para mascotas en seguimiento post-adopción (RF 4.1, RN 4.1)",
         "Mascota en proceso de adopción o seguimiento post-adopción.",
         "1. Acceder al módulo de seguimiento de la adopción.\n2. Redactar una recomendación/sugerencia de cuidado (ej. 'Mantener hidratación y cepillado semanal').\n3. Guardar el registro.",
         "El sistema almacena la sugerencia en la base de datos y la muestra en la línea de tiempo del seguimiento para consulta del adoptante.",
         "Media"),
        
        ("CP-AS-05", "Calificación y evaluación del nivel de cuidado de la mascota (RF 4.2, RN 4.2)",
         "Administrador o Refugio evaluando el estado de bienestar de la mascota en seguimiento.",
         "1. Abrir la ficha de seguimiento de la mascota.\n2. Asignar una calificación/puntuación al nivel de cuidado y bienestar.\n3. Guardar evaluación.",
         "El sistema registra la puntuación de cuidado y la asocia al historial de la adopción con fecha y evaluador.",
         "Media"),
        
        ("CP-AS-06", "Control de seguridad: bloqueo de acceso al panel /admin para usuarios no autorizados (RN 4.4, RNF 5)",
         "Usuario con rol Adoptante (USUARIO) o usuario no autenticado.",
         "1. Iniciar sesión como adoptante.\n2. Intentar ingresar manualmente a la URL 'http://localhost:3000/admin'.",
         "El sistema bloquea el acceso de inmediato, redirige al usuario a la página de login o inicio y la API responde con código HTTP 403 Forbidden.",
         "Crítica")
    ]
    
    for row in casos_m4:
        r = t8.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 10. Table 9: Registro de Ejecución (6 filas)
    t9 = doc.tables[9]
    while len(t9.rows) > 1:
        tr = t9.rows[-1]._tr
        t9._tbl.remove(tr)
        
    ejec_m4 = [
        ("CP-AS-01", "18/08/2026", "Kevin Ayala", "APROBADO", "Dashboard administrativo y métricas en tiempo real", "Ninguno"),
        ("CP-AS-02", "18/08/2026", "Kevin Ayala", "APROBADO", "Eliminación de cuenta confirmada y tabla actualizada", "Ninguno"),
        ("CP-AS-03", "18/08/2026", "Kevin Ayala", "APROBADO", "Refugio aprobado y rol actualizado en base de datos", "DEF-AS-01 (Corregido)"),
        ("CP-AS-04", "18/08/2026", "Kevin Ayala", "APROBADO", "Sugerencia de cuidado persistida en seguimiento", "Ninguno"),
        ("CP-AS-05", "18/08/2026", "Kevin Ayala", "APROBADO", "Calificación de bienestar registrada con éxito", "Ninguno"),
        ("CP-AS-06", "18/08/2026", "Kevin Ayala", "APROBADO", "Bloqueo estricto y redirección a rol adoptante", "Ninguno")
    ]
    for row in ejec_m4:
        r = t9.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 11. Table 11: Registro de Defectos
    t11 = doc.tables[11]
    while len(t11.rows) > 1:
        tr = t11.rows[-1]._tr
        t11._tbl.remove(tr)
        
    def_m4 = [
        ("DEF-AS-01", "Aprobación de refugio no actualizaba el rol del usuario a REFUGIO de forma inmediata",
         "1. Administrador aprueba solicitud de refugio en /admin/solicitudes-refugio. 2. Refugio inicia sesión.",
         "Esperado: Usuario obtiene rol REFUGIO y accede a su panel. Real: Mantenía rol USUARIO hasta segundo reinicio de sesión.",
         "Alta", "Cerrado", "Kevin Ayala")
    ]
    for row in def_m4:
        r = t11.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 12. Table 12: Riesgos y Contingencias
    t12 = doc.tables[12]
    while len(t12.rows) > 1:
        tr = t12.rows[-1]._tr
        t12._tbl.remove(tr)
        
    rsk_m4 = [
        ("RSK-01: Escalado de privilegios no autorizado hacia el rol ADMIN",
         "Baja / Crítica",
         "Validación estricta de roles en backend con Spring Security (@PreAuthorize y filtros en controladores).",
         "Kevin Ayala"),
        ("RSK-02: Eliminación accidental de un refugio con mascotas activas",
         "Baja / Alta",
         "Incorporar modales de confirmación de doble paso y validación de mascotas asociadas antes de borrar.",
         "Juan Alvarez"),
        ("RSK-03: Demora en la carga de métricas globales por crecimiento de usuarios",
         "Media / Media",
         "Uso de consultas agregadas COUNT directas en base de datos para responder en milisegundos.",
         "Kevin Ayala")
    ]
    for row in rsk_m4:
        r = t12.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 13. Table 13: Cronograma
    t13 = doc.tables[13]
    while len(t13.rows) > 1:
        tr = t13.rows[-1]._tr
        t13._tbl.remove(tr)
        
    cron_m4 = [
        ("Planificación y diseño de casos de prueba del Módulo 4", "Juan Alvarez", "17/08/2026", "18/08/2026", "Completado"),
        ("Configuración de cuentas de prueba (Admin, Refugio, Adoptante)", "Kevin Ayala", "18/08/2026", "18/08/2026", "Completado"),
        ("Ejecución de los casos de prueba CP-AS-01 al CP-AS-06", "Kevin Ayala", "18/08/2026", "18/08/2026", "Completado"),
        ("Verificación de seguridad y cierre del Módulo 4", "Kevin Ayala & Juan Alvarez", "18/08/2026", "18/08/2026", "Completado")
    ]
    for row in cron_m4:
        r = t13.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 14. Table 14: Cierre
    t14 = doc.tables[14]
    t14.rows[1].cells[1].text = "18/08/2026"
    t14.rows[3].cells[1].text = "Módulo 4: Administración del Sistema validado y aprobado en su totalidad. Cumple al 100% con los requisitos funcionales (RF 4.1, 4.2, 4.4, 4.5) y reglas de negocio asociadas."

    # 15. Table 15: Métricas de Resultados
    t15 = doc.tables[15]
    t15.rows[0].cells[1].text = "6"
    t15.rows[1].cells[1].text = "6"
    t15.rows[2].cells[1].text = "6 (100%)"
    t15.rows[3].cells[1].text = "0 (0%)"
    t15.rows[4].cells[1].text = "1"
    t15.rows[5].cells[1].text = "1 (100%)"
    t15.rows[6].cells[1].text = "0.60 segundos (Límite: <3s)"
    t15.rows[7].cells[1].text = "0.0% (Cero fallos en operaciones críticas)"
    t15.rows[8].cells[1].text = "100% de cobertura en RF 4.1, 4.2, 4.4, 4.5 y RN 4.1 a 4.5"

    # Save output docx with fallback if open
    try:
        doc.save(output_path)
        print("SUCCESS: File generated at", output_path)
    except PermissionError:
        alt_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_4_Administracion_del_Sistema_v2.docx'
        doc.save(alt_path)
        print("AVISO: Archivo bloqueado por Word. SUCCESS: Guardado como:", alt_path)

if __name__ == '__main__':
    generate_module_4_doc()
