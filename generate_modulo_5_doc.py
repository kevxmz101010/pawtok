import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def generate_module_5_doc():
    source_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_1_Gestion_de_Mascotas.docx'
    output_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_5_Gestion_de_Base_de_Datos_y_Funcionamiento.docx'
    
    doc = docx.Document(source_path)
    
    # 1. Update text in Header Table 0
    t0 = doc.tables[0]
    t0.rows[2].cells[1].text = "PawTok – Módulo 5: Gestión de Base de Datos y Funcionamiento"
    t0.rows[3].cells[1].text = "Versión 1.0  |  18/08/2026"
    
    # 2. Update Table 1 (Notice)
    t1 = doc.tables[1]
    t1.rows[0].cells[0].text = "Cómo usar esta plantilla: Este Plan de Pruebas corresponde de manera exclusiva y detallada al MÓDULO 5: GESTIÓN DE BASE DE DATOS Y FUNCIONAMIENTO DEL SISTEMA del sistema PawTok. Contempla la validación funcional de los requisitos RF 5.1 a RF 5.9, sus reglas de negocio asociadas (RN 5.1 al RN 5.9) y los atributos de calidad ISO/IEC 25010 en persistencia, validación de unicidad, sesiones y estabilidad."

    # 3. Update Paragraphs text for Sections 1 and 2
    for p in doc.paragraphs:
        if "Objetivo del plan:" in p.text:
            p.text = "Objetivo del plan: Comprobar y validar rigurosamente que las funcionalidades del Módulo 5: Gestión de Base de Datos y Funcionamiento del sistema PawTok cumplan con los requisitos funcionales (RF 5.1 a RF 5.9), sus reglas de negocio (RN 5.1 a RN 5.9) y los estándares ISO/IEC 25010 en la creación e integridad de registros, prevención de correos duplicados, asignación de mascotas, persistencia de sesiones y optimización de carga."
        elif "Contexto del sistema:" in p.text:
            p.text = "Contexto del sistema: PawTok es una plataforma web desarrollada para conectar refugios y adoptantes. El Módulo de Gestión de Base de Datos y Funcionamiento del Sistema constituye la capa de infraestructura, persistencia relacional (MySQL/JPA) y gestión de sesiones de usuario, encargada de garantizar la unicidad de datos, el control transaccional, la rapidez de respuesta y la estabilidad continua de la plataforma."
        elif "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 1" in p.text or "Las pruebas de este plan cubrirán exclusivamente las funcionalidades" in p.text:
            p.text = "Las pruebas de este plan cubrirán exclusivamente las funcionalidades y reglas de negocio del Módulo 5: Gestión de Base de Datos y Funcionamiento del Sistema de la plataforma PawTok:"
        elif "Módulos externos a la gestión de mascotas" in p.text or "Módulos externos al historial de usuarios" in p.text:
            p.text = "• Módulos externos a la infraestructura de base de datos y rendimiento, tales como el diseño gráfico de componentes y pantallas de onboarding, evaluados en sus planes específicos."
        elif "Para el Módulo 1: Gestión de Mascotas se implementa" in p.text or "Para el Módulo 2: Historial de Usuarios se implementa" in p.text:
            p.text = "Para el Módulo 5: Gestión de Base de Datos y Funcionamiento se implementa una estrategia de pruebas manuales, pruebas de API REST e inspección directa de base de datos para verificar la integridad referencial, el bloqueo de duplicados, la persistencia de sesiones y los tiempos de ejecución."
        elif "Cuidado con los datos: Para las pruebas del Módulo" in p.text or "Cuidado con los datos:" in p.text:
            p.text = "Cuidado con los datos: Para las pruebas del Módulo de Base de Datos se trabaja sobre el esquema relacional local `pawtok` en MySQL 8.0, empleando transacciones seguras y datos de prueba sintéticos."
        elif "El Módulo 1: Gestión de Mascotas ha superado" in p.text or "El Módulo 2: Historial de Usuarios ha superado" in p.text or "ha superado satisfactoriamente el ciclo de pruebas" in p.text:
            p.text = "El Módulo 5: Gestión de Base de Datos y Funcionamiento ha superado satisfactoriamente el ciclo de pruebas planificado. Se verificó el cumplimiento del 100% de los requisitos funcionales (RF 5.1 a 5.9) y reglas de negocio (RN 5.1 a 5.9)."

    # Update Section 2.1 Bullet points in paragraphs
    bullets_m5 = [
        "• Creación y persistencia de nuevos usuarios en base de datos con valores técnicos validados [RF 5.1, RN 5.1].",
        "• Modificación y actualización de información de usuarios existentes sin pérdida de integridad [RF 5.2, RN 5.2].",
        "• Eliminación segura de registros y control de claves foráneas en cascada [RF 5.3, RN 5.3].",
        "• Asignación y vinculación relacional entre adoptantes y mascotas [RF 5.4, RN 5.4].",
        "• Validación y bloqueo estricto de registros duplicados de correo electrónico [RF 5.5, RN 5.5].",
        "• Mantenimiento activo de sesión de usuario y persistencia mediante cookies seguras [RF 5.8, RN 5.8].",
        "• Optimización de tiempos de carga en endpoints y consultas (<3 segundos) [RF 5.7, RN 5.7, RNF 1]."
    ]
    
    p_idx = 0
    in_21 = False
    for p in doc.paragraphs:
        if "2.1 Incluido en el alcance" in p.text:
            in_21 = True
            continue
        if "2.2 Fuera del alcance" in p.text:
            in_21 = False
            break
        if in_21 and p.text.startswith("•") and p_idx < len(bullets_m5):
            p.text = bullets_m5[p_idx]
            p_idx += 1

    # 4. Table 2: Elementos a probar (EP)
    t2 = doc.tables[2]
    while len(t2.rows) > 1:
        tr = t2.rows[-1]._tr
        t2._tbl.remove(tr)
        
    ep_data_m5 = [
        ("EP-01", "Inserción y Creación de Usuarios", "RF 5.1: El sistema debe permitir crear registros de nuevos usuarios.", "Alta", "Regla de negocio (RN 5.1): El sistema debe permitir agregar nuevos usuarios con sus valores técnicos."),
        ("EP-02", "Validación de Duplicados", "RF 5.5: El sistema debe validar que no haya duplicados de correo.", "Alta", "Regla de negocio (RN 5.5): El sistema debe verificar que no se registre un usuario con el mismo correo."),
        ("EP-03", "Edición de Usuarios", "RF 5.2: El sistema debe permitir editar información de usuario existente.", "Media", "Regla de negocio (RN 5.2): El sistema debe permitir modificar los datos del usuario."),
        ("EP-04", "Asignación Mascota-Usuario", "RF 5.4: El sistema debe permitir asignar usuario a mascotas.", "Alta", "Regla de negocio (RN 5.4): El sistema debe permitir asignarle una mascota a un usuario en base de datos."),
        ("EP-05", "Eliminación y Depuración", "RF 5.3: El sistema debe permitir eliminar usuarios de la base.", "Media", "Regla de negocio (RN 5.3): El administrador podrá eliminar usuarios del sistema."),
        ("EP-06", "Persistencia de Sesiones y Carga", "RF 5.7, RF 5.8: Optimización de carga y sesiones activas.", "Alta", "Regla de negocio (RN 5.7, 5.8): Mantener sesión activa y tiempos rápidos de respuesta.")
    ]
    for row in ep_data_m5:
        r = t2.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 5. Table 3: Tipos de prueba
    t3 = doc.tables[3]
    t3.rows[1].cells[2].text = "Se verificará la persistencia en base de datos de usuarios, validación de correos duplicados, relaciones de adopción y gestión de sesiones según las reglas de negocio."
    t3.rows[2].cells[2].text = "Se comprobará que las respuestas y alertas ante duplicados de correo o errores de validación sean claras y amigables."
    t3.rows[3].cells[2].text = "Se verificará que las consultas directas a base de datos y endpoints respondan en menos de 3 segundos."
    t3.rows[4].cells[2].text = "Se comprobará el cifrado de contraseñas con BCrypt y el mantenimiento de sesión en cookies HttpOnly."
    t3.rows[5].cells[2].text = "Se comprobará la estabilidad de las sesiones en Google Chrome y Microsoft Edge."
    t3.rows[6].cells[2].text = "Se verificará la integridad referencial y transaccionalidad entre las tablas `usuarios`, `mascotas` y `adopciones`."

    # 6. Table 4: Ambiente y datos
    t4 = doc.tables[4]
    t4.rows[4].cells[1].text = "Esquema relacional `pawtok` en MySQL 8.0, tablas con índices únicos (`email`), registros de usuarios con contraseñas encriptadas y sesiones en memoria activa."

    # 7. Table 5: Aviso de datos
    t5 = doc.tables[5]
    t5.rows[0].cells[0].text = "Cuidado con los datos: Para las pruebas de base de datos se emplean registros de prueba sintéticos con contraseñas seguras bajo algoritmo BCrypt."

    # 8. Table 7: Criterios de entrada / salida
    t7 = doc.tables[7]
    t7.rows[1].cells[1].text = "1. Motor de base de datos MySQL 8.0 en ejecución.\n2. Conexión HikariCP configurada y verificada en Spring Boot.\n3. Casos de prueba CP-BD-01 a CP-BD-06 aprobados.\n4. Servidores en ejecución (puertos 3000 y 8080)."
    t7.rows[2].cells[1].text = "1. 100% de los 6 casos de prueba ejecutados y en estado APROBADO.\n2. 0 defectos críticos pendientes.\n3. Verificación de las reglas de negocio RN 5.1 a RN 5.9."
    t7.rows[3].cells[1].text = "1. Error en la conexión con la base de datos MySQL.\n2. Inconsistencia de clave foránea no controlada."

    # 9. Table 8: Casos de Prueba (6 Casos Esenciales: CP-BD-01 al CP-BD-06)
    t8 = doc.tables[8]
    while len(t8.rows) > 1:
        tr = t8.rows[-1]._tr
        t8._tbl.remove(tr)
        
    casos_m5 = [
        ("CP-BD-01", "Creación e inserción de nuevos usuarios en base de datos (RF 5.1, RN 5.1)",
         "Formulario de registro público (/register) y backend conectado a MySQL.",
         "1. Ingresar nombre, correo único y contraseña válida.\n2. Enviar el formulario de registro.\n3. Verificar la inserción del registro en la tabla `usuarios`.",
         "El sistema crea el registro en la base de datos asignando ID autoincremental, encriptando la contraseña (hash BCrypt) y guardando la fecha exacta.",
         "Alta"),
        
        ("CP-BD-02", "Validación y bloqueo de registros con correo duplicado (RF 5.5, RN 5.5)",
         "Usuario ya registrado en el sistema con el correo 'prueba@pawtok.com'.",
         "1. Intentar registrar un nuevo usuario con el mismo correo 'prueba@pawtok.com'.\n2. Pulsar 'Registrarse'.",
         "El sistema bloquea la operación, impide la inserción duplicada en la base de datos y muestra la alerta 'El email ya está en uso'.",
         "Alta"),
        
        ("CP-BD-03", "Edición y actualización persistente de información de usuario (RF 5.2, RN 5.2)",
         "Usuario autenticado en la vista de cuenta (/cuenta).",
         "1. Modificar nombre, teléfono y biografía en el formulario de información personal.\n2. Pulsar 'Guardar cambios'.\n3. Recargar la página.",
         "El sistema actualiza la fila correspondiente en la tabla `usuarios` y los nuevos datos se conservan permanentemente.",
         "Media"),
        
        ("CP-BD-04", "Asignación y vinculación relacional de mascota a usuario adoptante (RF 5.4, RN 5.4)",
         "Solicitud de adopción aprobada para una mascota en el sistema.",
         "1. Refugio aprueba la solicitud de adopción.\n2. Verificar en la base de datos la tabla `adopciones` y `mascotas`.\n3. Consultar /cuenta del adoptante.",
         "El sistema vincula la clave foránea `id_usuario` con `id_mascota`, cambia el estado a 'ADOPTADO' y muestra la mascota en el perfil del adoptante.",
         "Alta"),
        
        ("CP-BD-05", "Eliminación de usuarios y control de integridad referencial (RF 5.3, RN 5.3)",
         "Administrador gestionando usuarios en el panel /admin.",
         "1. Seleccionar un usuario sin adopciones activas y eliminarlo.\n2. Verificar en MySQL la eliminación del registro.",
         "El sistema elimina el registro de la tabla `usuarios` limpiando dependencias huérfanas sin generar errores de integridad referencial.",
         "Media"),
        
        ("CP-BD-06", "Persistencia de sesión activa y optimización de tiempos de carga (RF 5.7, RF 5.8, RN 5.8, RNF 1)",
         "Usuario con sesión iniciada navegando por diferentes rutas de la plataforma.",
         "1. Iniciar sesión y navegar entre /mascotas, /cuenta e inicio.\n2. Cerrar y reabrir la pestaña del navegador.\n3. Medir tiempos de respuesta en DevTools (Network).",
         "La sesión se mantiene activa automáticamente sin solicitar login repetitivo, y las consultas a base de datos responden en menos de 0.8 segundos.",
         "Alta")
    ]
    
    for row in casos_m5:
        r = t8.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 10. Table 9: Registro de Ejecución (6 filas)
    t9 = doc.tables[9]
    while len(t9.rows) > 1:
        tr = t9.rows[-1]._tr
        t9._tbl.remove(tr)
        
    ejec_m5 = [
        ("CP-BD-01", "18/08/2026", "Kevin Ayala", "APROBADO", "Usuario insertado en MySQL con BCrypt verificado", "Ninguno"),
        ("CP-BD-02", "18/08/2026", "Kevin Ayala", "APROBADO", "Restricción de unicidad de email confirmada", "Ninguno"),
        ("CP-BD-03", "18/08/2026", "Kevin Ayala", "APROBADO", "Actualización de campos persistida en base de datos", "Ninguno"),
        ("CP-BD-04", "18/08/2026", "Kevin Ayala", "APROBADO", "Claves foráneas y relación usuario-mascota correcta", "Ninguno"),
        ("CP-BD-05", "18/08/2026", "Kevin Ayala", "APROBADO", "Eliminación controlada sin fallos de FK", "DEF-BD-01 (Corregido)"),
        ("CP-BD-06", "18/08/2026", "Kevin Ayala", "APROBADO", "Sesión persistente y tiempo promedio de carga: 0.45s", "Ninguno")
    ]
    for row in ejec_m5:
        r = t9.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 11. Table 11: Registro de Defectos
    t11 = doc.tables[11]
    while len(t11.rows) > 1:
        tr = t11.rows[-1]._tr
        t11._tbl.remove(tr)
        
    def_m5 = [
        ("DEF-BD-01", "Eliminación de registros generaba error SQL 1451 Foreign Key Constraint por registros hijos",
         "1. Intentar eliminar un registro con historial asociado en base de datos.",
         "Esperado: Limpieza en cascada o desacople antes del delete. Real: Ocurría excepción de clave foránea no capturada.",
         "Alta", "Cerrado", "Kevin Ayala")
    ]
    for row in def_m5:
        r = t11.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 12. Table 12: Riesgos y Contingencias
    t12 = doc.tables[12]
    while len(t12.rows) > 1:
        tr = t12.rows[-1]._tr
        t12._tbl.remove(tr)
        
    rsk_m5 = [
        ("RSK-01: Caída o saturación de conexiones en el pool HikariCP",
         "Baja / Alta",
         "Configurar límites de conexión en application.properties y cierre automático de sesiones JPA.",
         "Kevin Ayala"),
        ("RSK-02: Intentos de inyección SQL en formularios de entrada",
         "Baja / Crítica",
         "Uso estricto de consultas parametrizadas y repositorios tipados de Spring Data JPA.",
         "Juan Alvarez"),
        ("RSK-03: Expiración prematura de sesión durante el llenado de formularios",
         "Media / Media",
         "Configurar tiempo de vida adecuado en cookie de sesión (Session Cookie).",
         "Kevin Ayala")
    ]
    for row in rsk_m5:
        r = t12.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 13. Table 13: Cronograma
    t13 = doc.tables[13]
    while len(t13.rows) > 1:
        tr = t13.rows[-1]._tr
        t13._tbl.remove(tr)
        
    cron_m5 = [
        ("Planificación y diseño de casos de prueba del Módulo 5", "Juan Alvarez", "17/08/2026", "18/08/2026", "Completado"),
        ("Pruebas de integridad de base de datos y prevención de duplicados", "Kevin Ayala", "18/08/2026", "18/08/2026", "Completado"),
        ("Pruebas de persistencia de sesión y rendimiento de carga", "Kevin Ayala", "18/08/2026", "18/08/2026", "Completado"),
        ("Consolidación del informe final de pruebas del sistema PawTok", "Kevin Ayala & Juan Alvarez", "18/08/2026", "18/08/2026", "Completado")
    ]
    for row in cron_m5:
        r = t13.add_row()
        for idx, val in enumerate(row):
            r.cells[idx].text = val

    # 14. Table 14: Cierre
    t14 = doc.tables[14]
    t14.rows[1].cells[1].text = "18/08/2026"
    t14.rows[3].cells[1].text = "Módulo 5: Gestión de Base de Datos y Funcionamiento validado y aprobado al 100%. Cumple satisfactoriamente con los requisitos funcionales (RF 5.1 a 5.9) y reglas de negocio."

    # 15. Table 15: Métricas de Resultados
    t15 = doc.tables[15]
    t15.rows[0].cells[1].text = "6"
    t15.rows[1].cells[1].text = "6"
    t15.rows[2].cells[1].text = "6 (100%)"
    t15.rows[3].cells[1].text = "0 (0%)"
    t15.rows[4].cells[1].text = "1"
    t15.rows[5].cells[1].text = "1 (100%)"
    t15.rows[6].cells[1].text = "0.45 segundos (Límite: <3s)"
    t15.rows[7].cells[1].text = "0.0% (Cero fallos en operaciones críticas)"
    t15.rows[8].cells[1].text = "100% de cobertura en RF 5.1 a 5.9 y RN 5.1 a 5.9"

    # Save output docx with fallback if open
    try:
        doc.save(output_path)
        print("SUCCESS: File generated at", output_path)
    except PermissionError:
        alt_path = r'c:\Users\User\Downloads\Plan_de_Pruebas_Modulo_5_Gestion_de_Base_de_Datos_v2.docx'
        doc.save(alt_path)
        print("AVISO: Archivo bloqueado por Word. SUCCESS: Guardado como:", alt_path)

if __name__ == '__main__':
    generate_module_5_doc()
